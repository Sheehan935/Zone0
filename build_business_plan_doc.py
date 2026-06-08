from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Zone 0 Landscaping - Paid Education Business Plan.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_run(run, size=None, bold=False, color="000000"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if not text:
        return p
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        style_run(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        style_run(rest)
    else:
        run = p.add_run(text)
        style_run(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        style_run(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        style_run(run)


def add_h1(doc, text):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=20)
    return p


def add_h2(doc, text):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=16)
    return p


def add_h3(doc, text):
    p = doc.add_heading(level=3)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=14, color="434343")
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], "F8F9FA")
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        style_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            style_run(run)
    set_table_width(table, widths)
    set_table_borders(table)
    doc.add_paragraph()
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.color.rgb = RGBColor.from_string("000000")


def build():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.line_spacing = 1.15
    title_run = title.add_run("Zone 0 Landscaping Paid Education Business Plan")
    style_run(title_run, size=26)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("Working plan for zone0landscaping.com | Google Docs-ready draft")
    style_run(run, color="555555")

    add_paragraph(
        doc,
        "This plan starts Zone 0 Landscaping as a paid educational site before expanding into remote reviews, on-site inspections, HOA workshops, and contractor-coordinated mitigation work. The strategy is to sell clarity first, then convert the most motivated homeowners into higher-touch services.",
    )

    add_h1(doc, "1. Executive Summary")
    add_paragraph(
        doc,
        "Zone 0 Landscaping will launch as a California-focused wildfire-readiness education business for homeowners in fire-risk areas. The first paid product is a practical Zone 0 homeowner guide supported by a free self-check, photo checklist, and email funnel.",
    )
    add_paragraph(
        doc,
        "The business should avoid leading with landscaping or AI. The clearer promise is: Learn what to fix before you hire anyone. This positions the company as a trusted educator and consultant, not just another contractor.",
    )

    add_h1(doc, "2. Core Positioning")
    add_paragraph(doc, "Category: Zone 0 wildfire-readiness education and consulting.")
    add_paragraph(doc, "Primary promise: Know what to fix. Get the work done. Keep the proof.")
    add_paragraph(doc, "Homepage message: Understand your Zone 0 risks before fire season.")
    add_paragraph(
        doc,
        "Subheadline: A practical California homeowner guide to the first five feet around your home: what to remove, what to replace, how to document the work, and when to call a professional.",
    )

    add_h1(doc, "3. Offer Ladder")
    add_table(
        doc,
        ["Tier", "Offer", "Price", "Purpose"],
        [
            ["Free", "Zone 0 self-check quiz and photo checklist", "$0", "Capture local leads and identify urgency."],
            ["Entry", "California Zone 0 Homeowner Readiness Guide", "$29-$49", "Validate paid demand and establish authority."],
            ["Kit", "DIY Zone 0 Action Plan Kit", "$99", "Provide templates, documentation packet, tracker, and contractor scope worksheet."],
            ["Review", "Remote Photo Review", "$199-$299", "Convert motivated homeowners into personalized guidance."],
            ["Consulting", "On-site Zone 0 Inspection", "$399-$799", "Walkthrough, photos, priority action plan, and documentation packet."],
            ["B2B", "HOA or neighborhood workshop", "$1,500+", "Acquire multiple homeowners through one trusted channel."],
            ["Service", "Clearing, hardscape, and contractor coordination", "Custom", "Add only after reliable crews and scope control are in place."],
        ],
        [0.9, 2.2, 1.0, 2.4],
    )

    add_h1(doc, "4. Target Customers")
    add_paragraph(doc, "Start with homeowners who already feel pressure from fire risk, insurance, HOA guidance, real estate transactions, or confusing local requirements.")
    add_bullets(
        doc,
        [
            "East Bay hills: Oakland Hills, Berkeley Hills, Orinda, Lafayette, Moraga, Walnut Creek, Alamo, and Danville.",
            "North Bay and coastal risk markets: Marin, Sonoma, Napa, Los Gatos, and the Santa Cruz Mountains.",
            "Homeowners with insurance anxiety, FAIR Plan exposure, or recent non-renewal concerns.",
            "HOA board members and neighborhood leaders who can create bulk demand.",
            "Real estate agents and insurance brokers who need practical referral resources.",
        ],
    )

    add_h1(doc, "5. Paid Guide Product")
    add_h2(doc, "Product Name")
    add_paragraph(doc, "The California Zone 0 Homeowner Readiness Guide")
    add_h2(doc, "What It Includes")
    add_bullets(
        doc,
        [
            "Plain-English Zone 0 explainer.",
            "The bubble of starvation mental model.",
            "0-5 foot photo checklist.",
            "Primary ignition pathway checklist.",
            "Severity versus priority matrix.",
            "Red Flag event protocol.",
            "Insurance and grant documentation checklist.",
            "Maintenance calendar.",
            "Contractor scope worksheet.",
            "Local authority and compliance caveats.",
        ],
    )
    add_paragraph(
        doc,
        "The guide should feel like a field manual and homeowner workbook, not a generic ebook. The buyer should finish with a specific list of photos to take, items to remove, questions to ask, and documentation to keep.",
    )

    add_h1(doc, "6. Website Structure")
    add_numbered(
        doc,
        [
            "Homepage: explain the Zone 0 risk, introduce the free self-check, and present the paid guide.",
            "Free Self-Check: capture email, ZIP code, insurance concern, HOA status, and optional photo-upload interest.",
            "Paid Guide Sales Page: describe the guide, preview the table of contents, and connect the purchase to practical homeowner clarity.",
            "Thank-You or Member Page: deliver the guide and upsell remote photo review.",
            "Photo Review Page: explain what photos to submit and what the homeowner receives.",
            "Resources Section: publish local and high-intent educational articles.",
        ],
    )

    add_h1(doc, "7. Content Strategy")
    add_paragraph(doc, "Build authority through local, high-intent pages that answer specific homeowner searches.")
    add_bullets(
        doc,
        [
            "California Zone 0 checklist.",
            "Zone 0 vs defensible space.",
            "Zone 0 mulch rules.",
            "Wood fence-to-home wildfire risk.",
            "Wildfire insurance documentation checklist.",
            "East Bay Zone 0 homeowner guide.",
            "Oakland Hills wildfire-readiness checklist.",
            "Red Flag event homeowner checklist.",
            "What to photograph for wildfire insurance records.",
        ],
    )

    add_h1(doc, "8. Launch Plan")
    add_table(
        doc,
        ["Phase", "Timeline", "Actions", "Success Metric"],
        [
            ["Build", "Week 1", "Create paid guide, free checklist, sales page, Stripe link, and email capture.", "Site can take payment and deliver guide."],
            ["Pilot", "Week 2", "Offer discounted guide and remote reviews to local homeowners, agents, and brokers.", "20 paid guide sales and 5 reviews."],
            ["Local Authority", "Weeks 3-4", "Publish 4 local articles and collect questions from buyers.", "100 email subscribers."],
            ["Partnership", "Month 2", "Pitch HOA talks, real estate offices, and insurance brokers.", "2 workshops booked."],
            ["Service Expansion", "Month 3+", "Add on-site inspections and contractor coordination after workflow is proven.", "5 inspections per month."],
        ],
        [0.9, 0.9, 3.2, 1.5],
    )

    add_h1(doc, "9. Revenue Targets")
    add_table(
        doc,
        ["Milestone", "Target Activity", "Estimated Revenue"],
        [
            ["Month 1", "100 subscribers, 20 guide sales, 5 remote reviews", "$1,000-$2,000"],
            ["Month 3", "500 subscribers, 100 guide sales, 20 reviews, 5 inspections", "$6,000-$12,000"],
            ["Month 6", "300 guide sales, 50 reviews, 20 inspections, 2 HOA workshops", "$25,000-$50,000 cumulative"],
        ],
        [1.1, 3.8, 1.6],
    )

    add_h1(doc, "10. Risk Controls")
    add_paragraph(doc, "The education-first model reduces operational complexity, but the site still needs careful language.")
    add_bullets(
        doc,
        [
            "Use wildfire-readiness, documentation, and educational review language.",
            "Avoid fireproof, guaranteed compliant, CAL FIRE certified, or insurance discount guaranteed claims.",
            "State that formal determinations depend on local authorities, insurers, grant programs, or inspectors.",
            "Do not recommend structural work without licensed professionals.",
            "Keep the paid guide current as California and local guidance evolves.",
        ],
    )

    add_h1(doc, "11. Immediate Next Steps")
    add_numbered(
        doc,
        [
            "Turn the current homepage into an education-first landing page with the paid guide as the main offer.",
            "Create the first paid guide PDF or member page from the consultant master guide.",
            "Add Stripe payment and a simple fulfillment flow.",
            "Create a free Zone 0 self-check that captures email, ZIP, insurance concern, and HOA status.",
            "Offer the first 10 remote photo reviews manually to learn what homeowners actually ask.",
        ],
    )

    add_h1(doc, "12. Working Brand Guardrails")
    add_paragraph(doc, "Sound like a sober field consultant, not a fear marketer or generic landscaper.")
    add_bullets(
        doc,
        [
            "Use: ember-resistant buffer, ignition pathways, action plan, documentation, maintenance protocol.",
            "Avoid: guaranteed compliance, fireproof, certified safe, AI-powered landscaping.",
            "Lead with clarity and practical action before selling labor.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
